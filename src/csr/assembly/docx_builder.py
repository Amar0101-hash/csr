"""Render generated sections into the template while preserving its structure.

Strategy: load the template, walk it in document order, and for each section
that we authored, replace the colored *guidance* paragraphs in place with the
generated prose — keeping every heading, boilerplate paragraph, and template
table exactly where the template put them. Sections we did not author (title
page, signature page, annex lists) are left untouched so the template's
authoring instructions remain visible to a human reviewer.

Traceability is attached as Word comments anchored to each generated paragraph.
"""
from __future__ import annotations

from collections import Counter
from pathlib import Path
from typing import Optional

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..config import GUIDANCE_COLORS
from ..ingestion.docx_reader import iter_block_items
from ..ingestion.template_parser import HEADING_STYLES, SKIP_STYLES
from ..models import GeneratedSection, SectionSpec

GUIDANCE_COLOR_SET = set(GUIDANCE_COLORS.keys())
BODY_STYLE = "Document Text"


def _dominant_color(p: Paragraph) -> Optional[str]:
    counts: Counter = Counter()
    for r in p.runs:
        if r.text.strip():
            col = None
            try:
                if r.font.color is not None and r.font.color.rgb is not None:
                    col = str(r.font.color.rgb).upper()
            except Exception:
                col = None
            counts[col] += len(r.text)
    if not counts:
        return None
    return counts.most_common(1)[0][0]


def _delete_paragraph(p: Paragraph) -> None:
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _safe_style(doc: Document, name: str) -> Optional[str]:
    try:
        _ = doc.styles[name]
        return name
    except KeyError:
        return None


def build_report(
    template_path: Path,
    sections: list[SectionSpec],
    generated: dict[str, GeneratedSection],
    output_path: Path,
    add_comments: bool = True,
) -> dict:
    doc = Document(str(template_path))
    body_style = _safe_style(doc, BODY_STYLE)

    # First pass: assign each block to a section index (aligned by heading order
    # with the parsed `sections`). Collect guidance paragraphs, the heading
    # paragraph, and the template tables per section.
    sec_idx = -1
    guidance_by_section: dict[int, list[Paragraph]] = {}
    heading_by_section: dict[int, Paragraph] = {}
    tables_by_section: dict[int, list[Table]] = {}

    for block in iter_block_items(doc):
        if isinstance(block, Table):
            if sec_idx >= 0:
                tables_by_section.setdefault(sec_idx, []).append(block)
            continue
        p = block
        style = p.style.name if p.style else "Normal"
        text = p.text.strip()

        is_heading = style in HEADING_STYLES and (text or style == "Title")
        if is_heading:
            sec_idx += 1
            heading_by_section[sec_idx] = p
            continue
        if sec_idx < 0 or style in SKIP_STYLES or not text:
            continue

        color = _dominant_color(p)
        if color in GUIDANCE_COLOR_SET and color != "0000FF":
            guidance_by_section.setdefault(sec_idx, []).append(p)

    # Second pass: fill tables, rename placeholder headings, insert prose, remove
    # guidance. Iterate a stable snapshot so edits don't disturb traversal.
    stats = {"authored": 0, "placeholder": 0, "skipped": 0, "comments": 0, "tables_filled": 0}

    for idx, section in enumerate(sections):
        gen = generated.get(section.key)
        g_paras = guidance_by_section.get(idx, [])
        sec_tables = tables_by_section.get(idx, [])

        if gen is None:
            if not section.generate:
                stats["skipped"] += 1
                continue
            # generate=True but nothing produced -> placeholder, under the heading
            ins = _Inserter(
                doc,
                before=g_paras[0] if g_paras else None,
                after_el=heading_by_section[idx]._p if (not g_paras and idx in heading_by_section) else None,
            )
            ins.paragraph(_placeholder_text(None), body_style)
            stats["placeholder"] += 1
            for gp in g_paras:
                _delete_paragraph(gp)
            continue

        # (a) rename placeholder endpoint heading, if the model gave a real name
        if gen.heading_override and idx in heading_by_section:
            _set_heading_text(heading_by_section[idx], gen.heading_override)

        # (b) fill the section's template form-tables
        for tf in gen.table_fills:
            if 0 <= tf.table_index < len(sec_tables):
                if _fill_form_table(sec_tables[tf.table_index], tf):
                    stats["tables_filled"] += 1

        # (c) insert prose under the heading. Anchor before the first guidance
        # paragraph if any; otherwise insert right after the heading so sections
        # without guidance (endpoint placeholders) aren't dumped at doc end.
        ins = _Inserter(
            doc,
            before=g_paras[0] if g_paras else None,
            after_el=heading_by_section[idx]._p if (not g_paras and idx in heading_by_section) else None,
        )
        first_para: Optional[Paragraph] = None
        if gen.paragraphs:
            for para_text in gen.paragraphs:
                if _looks_like_md_table(para_text):
                    ins.table(_parse_md_table(para_text))
                else:
                    np = ins.paragraph(para_text, body_style)
                    if first_para is None:
                        first_para = np
            if add_comments and first_para is not None:
                _attach_comment(doc, first_para, gen, stats)
            stats["authored"] += 1
        elif not gen.table_fills:
            ins.paragraph(_placeholder_text(gen), body_style)
            stats["placeholder"] += 1
        else:
            stats["authored"] += 1  # table-only section (Title Page / Summary)

        for gp in g_paras:
            _delete_paragraph(gp)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    written = _save_resilient(doc, output_path)
    stats["output_path"] = str(written)
    return stats


class _Inserter:
    """Places paragraphs/tables at the right spot: before the guidance anchor if
    the section has guidance, otherwise chained right after the heading. This
    keeps generated content under its heading instead of at the document end."""

    def __init__(self, doc: Document, before: Optional[Paragraph] = None, after_el=None):
        self.doc = doc
        self.before = before
        self.cursor = after_el  # lxml element to addnext against (heading/last block)

    def _place(self, new_el) -> None:
        if self.before is not None:
            self.before._p.addprevious(new_el)
        elif self.cursor is not None:
            self.cursor.addnext(new_el)
            self.cursor = new_el
        # else: leave where doc.add_* put it (doc end) as a last resort

    def paragraph(self, text: str, style) -> Paragraph:
        p = self.doc.add_paragraph(text, style=style)
        self._place(p._p)
        return p

    def table(self, rows: list[list[str]]) -> None:
        tbl = _build_table(self.doc, rows)
        if tbl is not None:
            self._place(tbl._tbl)


def _insert_note(doc: Document, anchor: Optional[Paragraph], text: str, style) -> Paragraph:
    if anchor is not None:
        return anchor.insert_paragraph_before(text, style=style)
    return doc.add_paragraph(text, style=style)


def _set_heading_text(heading: Paragraph, text: str) -> None:
    """Replace a placeholder heading's text with a single fresh run so it inherits
    the Heading style's formatting (bold, size, black) instead of the placeholder
    runs' inconsistent per-run overrides (blue, mixed weights)."""
    for r in list(heading.runs):
        r._element.getparent().remove(r._element)
    heading.add_run(text)  # fresh run, no overrides -> inherits Heading style


# ---- markdown table rendering ----

def _looks_like_md_table(text: str) -> bool:
    lines = [l for l in text.splitlines() if l.strip()]
    piped = [l for l in lines if l.strip().startswith("|")]
    return len(piped) >= 2


def _parse_md_table(text: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in text.splitlines():
        s = line.strip()
        if not s.startswith("|"):
            continue
        body = s.strip("|")
        if set(body.replace("|", "").strip()) <= set("-: "):
            continue  # separator row
        rows.append([c.strip() for c in body.split("|")])
    return rows


def _build_table(doc: Document, rows: list[list[str]]) -> Optional[Table]:
    """Build a Word table (bold header row) appended at doc end; the caller moves
    it into place."""
    if not rows:
        return None
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for i, r in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.text = r[j] if j < len(r) else ""
            if i == 0:
                for cp in cell.paragraphs:
                    for run in cp.runs:
                        run.bold = True
    return table


# ---- template form-table filling ----

def _snapshot(run) -> tuple:
    """Capture (bold, size, name, color) of a run before we clear it, so the
    template's own label formatting can be reapplied (preserve fonts/styles)."""
    if run is None:
        return (None, None, None, None)
    color = None
    try:
        if run.font.color is not None and run.font.color.rgb is not None:
            color = run.font.color.rgb
    except Exception:
        color = None
    return (run.bold, run.font.size, run.font.name, color)


def _apply(run, snap: tuple, bold, black: bool) -> None:
    b, size, name, color = snap
    run.bold = b if bold is None else bold
    try:
        if size is not None:
            run.font.size = size
        if name is not None:
            run.font.name = name
        if black:
            run.font.color.rgb = None            # body value -> default (black)
        elif color is not None:
            run.font.color.rgb = color           # preserve template label color
    except Exception:
        pass


def _first_run(cell) -> object:
    for p in cell.paragraphs:
        if p.runs:
            return p.runs[0]
    return None


def _write_cell(cell, label: Optional[str], value: str) -> None:
    """Rewrite a cell preserving the template's label formatting.

    synopsis  (label given): keep the original label run's formatting exactly
              (bold/size/name/color as the template had it), value in body style.
    label_value (label None): this is the value cell; write value in body style.
    """
    para = cell.paragraphs[0]
    orig = para.runs[0] if para.runs else _first_run(cell)
    snap = _snapshot(orig)
    for p in cell.paragraphs[1:]:
        _delete_paragraph(p)
    for r in list(para.runs):
        r.text = ""
    first = para.runs[0] if para.runs else para.add_run("")
    if label:
        first.text = f"{label}: "
        _apply(first, snap, bold=snap[0], black=False)   # preserve template label
        val_run = para.add_run(value)
        _apply(val_run, snap, bold=False, black=True)    # body value
    else:
        first.text = value
        _apply(first, snap, bold=False, black=True)       # body value


def _fill_form_table(table: Table, tf) -> bool:
    from ..generation.table_fill import _clean_label

    filled = False
    if tf.mode == "synopsis":
        for row in table.rows:
            cell = row.cells[0]
            cell_text = cell.text.strip()
            if ":" not in cell_text:
                continue
            label = _clean_label(cell_text)
            val = tf.values.get(label)
            if val:
                _write_cell(cell, label, val)
                filled = True
    else:  # label_value
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = _clean_label(cells[0].text)
            val = tf.values.get(label)
            if val:
                _write_cell(cells[1], None, val)
                filled = True
    return filled


def _save_resilient(doc: Document, output_path: Path) -> Path:
    """Save the docx. If the target is locked (open in Word/IDE), fall back to a
    timestamped filename instead of losing the whole run."""
    try:
        doc.save(str(output_path))
        return output_path
    except PermissionError:
        import time

        stamp = time.strftime("%Y%m%d_%H%M%S")
        alt = output_path.with_name(f"{output_path.stem}_{stamp}{output_path.suffix}")
        doc.save(str(alt))
        print(
            f"[assemble] '{output_path.name}' is open/locked — wrote '{alt.name}' instead. "
            f"Close it and re-run to update the primary file.",
            flush=True,
        )
        return alt


def _placeholder_text(gen: Optional[GeneratedSection]) -> str:
    """Shown when a section produced no content, so it's never left blank."""
    note = (gen.notes if gen and gen.notes else "").strip()
    base = "Data not available in the provided source documents for this section."
    return f"{base} [{note}]" if note else base


def _attach_comment(doc: Document, para: Paragraph, gen: GeneratedSection, stats: dict) -> None:
    lines = ["Generated from study sources."]
    if gen.citations:
        srcs = []
        for c in gen.citations[:8]:
            srcs.append(f"{c.doc}: {c.section_path}")
        # de-dup preserving order
        seen = set()
        uniq = [s for s in srcs if not (s in seen or seen.add(s))]
        lines.append("Sources: " + "; ".join(uniq))
    v = gen.verification or {}
    if v.get("unsupported_numbers"):
        lines.append("REVIEW numbers not found verbatim in sources: " + ", ".join(v["unsupported_numbers"][:10]))
    if gen.notes:
        lines.append("Note: " + gen.notes)
    try:
        runs = [r for r in para.runs if r.text.strip()] or para.runs
        if runs:
            doc.add_comment(runs=runs, text="\n".join(lines), author="CSR-GenAI", initials="AI")
            stats["comments"] += 1
    except Exception:
        pass
